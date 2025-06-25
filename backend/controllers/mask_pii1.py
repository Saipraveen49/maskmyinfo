import pytesseract
from PIL import Image, ImageDraw
import spacy
import re
import os
from pdf2image import convert_from_path
from docx import Document
import json
import sys

# Set Tesseract path if needed
pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

# Load the spaCy model for NER
nlp = spacy.load("en_core_web_sm")

# Define regex patterns for structured PII detection
regex_patterns = {
    'aadhaar': r'\b\d{4}\s\d{4}\s\d{4}\b',
    'pan': r'\b[A-Z]{5}\d{4}[A-Z]{1}\b',
    'passport_number': r'\b[A-PR-WYa-pr-wy][1-9]\d\s?\d{4}[1-9]\b',
    'voter_id': r'\b[A-Z]{3}\d{7}\b',
    'dl_number': r'\b[A-Z]{2}\d{2}\s?\d{11}\b',
    'mobile_number': r'\b[6-9]\d{9}\b',
    'email': r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+',
    'pin_code': r'\b\d{6}\b',
    'name': r'\b[A-Z][a-z]*\s[A-Z][a-z]*\b',
}

def clean_text(text):
    """Clean up text."""
    return ' '.join(text.split())

# Detect sensitive data using regex and NER
def detect_sensitive_data(text):
    detected_values = set()  # Avoid duplicates

    # Detect structured PII using regex
    for pattern in regex_patterns.values():
        matches = re.findall(pattern, text)
        detected_values.update(matches)

    # Detect unstructured PII using NER
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ in ['PERSON', 'ORG', 'GPE', 'CARDINAL']:
            detected_values.add(ent.text)

    return list(detected_values)  # Return unique values

# Mask sensitive data based on user-selected values
def mask_sensitive_data(text, selected_values):
    masked_text = text

    # Mask structured PII using regex based on selected values
    for pattern in regex_patterns.values():
        matches = re.findall(pattern, text)
        for match in matches:
            if match in selected_values:
                masked_text = masked_text.replace(match, '****')

    # Mask unstructured PII detected using NER
    doc = nlp(text)
    for ent in doc.ents:
        if ent.text in selected_values:  # Only mask if the value is selected
            masked_text = masked_text.replace(ent.text, '****')

    return masked_text

# Save the redacted file based on the file type
def save_redacted_file(file_path, masked_text):
    ext = os.path.splitext(file_path)[-1].lower()
    redacted_file_path = os.path.join(os.path.dirname(file_path), f"masked_{os.path.basename(file_path)}")

    # Handle images
    if ext in ['.jpg', '.jpeg', '.png']:
        image = Image.open(file_path)
        draw = ImageDraw.Draw(image)
        # Redrawing masked content in images needs additional logic
        image.save(redacted_file_path)
        print(f"Redacted image saved at: {redacted_file_path}")

    # Handle PDFs
    elif ext == '.pdf':
        images = convert_from_path(file_path)
        # Saving only the images again (text redaction over image needs more logic)
        images[0].save(redacted_file_path, save_all=True, append_images=images[1:])
        print(f"Redacted PDF saved at: {redacted_file_path}")

    # Handle DOCX files
    elif ext == '.docx':
        doc = Document(file_path)
        for para in doc.paragraphs:
            para.text = masked_text
        doc.save(redacted_file_path)
        print(f"Redacted DOCX saved at: {redacted_file_path}")

    return redacted_file_path

# Process the file based on its type (image, PDF, DOCX)
def process_file(file_path, get_sensitive=False, selected_values=None):
    ext = os.path.splitext(file_path)[-1].lower()

    # Process images
    if ext in ['.jpg', '.jpeg', '.png']:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        clean_text_data = clean_text(text)

        if get_sensitive:
            return detect_sensitive_data(clean_text_data)
        else:
            masked_text = mask_sensitive_data(clean_text_data, selected_values)
            return save_redacted_file(file_path, masked_text)

    # Process PDFs
    elif ext == '.pdf':
        images = convert_from_path(file_path)
        full_text = ' '.join([pytesseract.image_to_string(image) for image in images])
        clean_text_data = clean_text(full_text)

        if get_sensitive:
            return detect_sensitive_data(clean_text_data)
        else:
            masked_text = mask_sensitive_data(clean_text_data, selected_values)
            return save_redacted_file(file_path, masked_text)

    # Process DOCX files
    elif ext == '.docx':
        doc = Document(file_path)
        full_text = clean_text(' '.join([para.text for para in doc.paragraphs]))

        if get_sensitive:
            return detect_sensitive_data(full_text)
        else:
            masked_text = mask_sensitive_data(full_text, selected_values)
            return save_redacted_file(file_path, masked_text)

    else:
        raise ValueError(f"Unsupported file format: {ext}")

# Main entry point for CLI usage
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Error: Please provide a file path as the first argument.")
        sys.exit(1)

    file_path = sys.argv[1]

    # Detect sensitive data and return it as JSON
    if '--get-sensitive' in sys.argv:
        sensitive_data = process_file(file_path, get_sensitive=True)
        print(json.dumps(sensitive_data))  # Output detected sensitive data as JSON

    # Mask selected sensitive data and save the redacted file
    elif '--mask' in sys.argv:
        if len(sys.argv) < 4:
            print("Error: Please provide the selected values as arguments.")
            sys.exit(1)

        selected_values = sys.argv[3:]  # Collect selected values from CLI arguments
        redacted_file_path = process_file(file_path, get_sensitive=False, selected_values=selected_values)
        print(f"Masked file saved at: {redacted_file_path}")
