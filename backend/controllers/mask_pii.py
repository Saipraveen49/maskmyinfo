import pytesseract
from PIL import Image, ImageDraw
import spacy
import re
import io
from pdf2image import convert_from_bytes
from docx import Document
import sys
import os

# Set Tesseract path
pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

# Load the spaCy model for NER
nlp = spacy.load("en_core_web_sm")

# Comprehensive regex patterns for structured PII
regex_patterns = {
    # Define your regex patterns here as needed
     'aadhaar': r'\b\d{4}\s\d{4}\s\d{4}\b',
    'pan': r'\b[A-Z]{5}\d{4}[A-Z]{1}\b',
    'passport_number': r'\b[A-PR-WYa-pr-wy][1-9]\d\s?\d{4}[1-9]\b',
    'voter_id': r'\b[A-Z]{3}\d{7}\b',
    'dl_number': r'\b[A-Z]{2}\d{2}\s?\d{11}\b',
    'ration_card': r'\b\d{2}[A-Z]{1}\d{2}\s?\d{5}\b',
    'pmjay': r'\b\d{4}\s\d{4}\s\d{4}\s\d{4}\b',
    'pmkisan': r'\b\d{12}\b',
    'uan': r'\b\d{12}\b',
    'epfo': r'\b[A-Z]{3}\d{7}\b',
    'esic': r'\b\d{10}\b',
    'gstin': r'\b\d{2}[A-Z]{5}\d{4}[A-Z]{1}[A-Z\d]{1}[Z]{1}[A-Z\d]{1}\b',
    'ifsc': r'\b[A-Z]{4}0[A-Z0-9]{6}\b',
    'tan': r'\b[A-Z]{4}\d{5}[A-Z]{1}\b',
    'cin': r'\b[U|L]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\b',
    'nsdl_pan_token': r'\b[A-Z]{3}\d{5}[A-Z]\b',
    'bank_account': r'\b\d{9,18}\b',
    'credit_card': r'\b(?:\d[ -]*?){13,16}\b',
    'mobile_number': r'\b[6-9]\d{9}\b',
    'email': r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+',
    'pin_code': r'\b\d{6}\b',
    'ip_address': r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b',
    'mac_address': r'\b([0-9A-Fa-f]{2}[:-]){5}([0-9A-Fa-f]{2})\b',
    'address': r'\b\d{1,3}\s\w+\s(street|road|lane|avenue|ave|rd|ln|st|blvd|drive|dr|plaza|plz|circle|cir|square|sq|parkway|pkwy|court|ct|hwy|expressway|expwy)\b',
    'birthdate': r'\b(?:0?[1-9]|[12][0-9]|3[01])[-/](?:0?[1-9]|1[012])[-/](?:19|20)?\d{2}\b',
    'vehicle_registration': r'\b[A-Z]{2}\s?\d{1,2}\s?[A-Z]{1,2}\s?\d{4}\b',
    'insurance_policy': r'\b[A-Z0-9]{5,12}\b',
    'age': r'\b(?:age|dob|born)\s*(\d{1,3})\b',
    'name': r'\b[A-Z][a-z]*\s[A-Z][a-z]*\b',
    'personal_info': r'\b(name|address|phone|mobile|email|dob|age|occupation|salary|income|gender)\b',
}

def detect_and_redact_text(text):
    sensitive_data = []
    masked_text = text

    # Detect structured PII using regex
    for label, pattern in regex_patterns.items():
        matches = re.findall(pattern, text)
        for match in matches:
            sensitive_data.append((match, label))
            masked_text = masked_text.replace(match, '[REDACTED]')

    # Detect entities using NER
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ in ['PERSON', 'ORG', 'GPE', 'CARDINAL']:
            sensitive_data.append((ent.text, ent.label_))
            masked_text = masked_text.replace(ent.text, '[REDACTED]')

    return masked_text, sensitive_data

def process_image(image_data):
    image = Image.open(io.BytesIO(image_data))
    data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)

    texts = data['text']
    bbox = [(data['left'][i], data['top'][i], data['width'][i], data['height'][i]) for i in range(len(texts))]

    full_text = ' '.join(texts)
    masked_text, sensitive_data = detect_and_redact_text(full_text)

    print("Detected Sensitive Data:", sensitive_data, file=sys.stderr)

    draw = ImageDraw.Draw(image)
    for (x, y, w, h), text in zip(bbox, texts):
        if text in full_text and text not in masked_text:
            draw.rectangle([x, y, x + w, y + h], fill="black")

    output = io.BytesIO()
    image.save(output, format='PNG')
    output.seek(0)
    return output.getvalue()

def process_pdf(pdf_data):
    try:
        # Convert PDF pages to images
        images = convert_from_bytes(pdf_data)
    except Exception as e:
        raise ValueError(f"Error converting PDF to images: {e}")

    redacted_images = []

    for page_num, image in enumerate(images):
        try:
            # Convert image to bytes
            image_bytes = io.BytesIO()
            image.save(image_bytes, format='PNG')
            image_bytes = image_bytes.getvalue()

            # Process the image
            redacted_image = process_image(image_bytes)
            redacted_images.append(Image.open(io.BytesIO(redacted_image)))
        except Exception as e:
            print(f"Error processing page {page_num + 1} of PDF: {e}", file=sys.stderr)
            continue

    if not redacted_images:
        raise ValueError("No pages could be processed from the PDF.")

    output = io.BytesIO()
    redacted_images[0].save(output, save_all=True, append_images=redacted_images[1:], format='PDF')
    output.seek(0)
    return output.getvalue()


def process_docx(docx_data):
    doc = Document(io.BytesIO(docx_data))
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    full_text = ' '.join(full_text)
    masked_text, sensitive_data = detect_and_redact_text(full_text)

    print("Detected Sensitive Data:", sensitive_data, file=sys.stderr)

    for para in doc.paragraphs:
        if para.text in full_text:
            para.text = para.text.replace(para.text, masked_text)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

def process_file(file_data, file_ext):
    if file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.webp']:
        return process_image(file_data)
    elif file_ext == '.pdf':
        return process_pdf(file_data)
    elif file_ext == '.docx':
        return process_docx(file_data)
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python script.py <input_file> <output_file>", file=sys.stderr)
        sys.exit(1)

    input_file_path = sys.argv[1]
    output_file_path = sys.argv[2]

    try:
        with open(input_file_path, 'rb') as f:
            file_data = f.read()

        file_ext = os.path.splitext(input_file_path)[-1].lower()
        result_data = process_file(file_data, file_ext)

        with open(output_file_path, 'wb') as f:
            f.write(result_data)

    except Exception as e:
        print(f"Error processing file: {e}", file=sys.stderr)
        sys.exit(1)
